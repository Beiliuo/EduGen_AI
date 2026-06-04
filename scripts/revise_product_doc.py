from pathlib import Path

from docx import Document
from docx.shared import RGBColor


INPUT = Path(r"C:\Users\Liu\Desktop\WK09-课程项目产品设计模板(2).docx")
OUTPUT = Path(r"C:\Users\Liu\Desktop\Code\WK09-课程项目产品设计模板_EduGenAI修订版.docx")

RED = RGBColor(220, 38, 38)


def add_run(paragraph, text, *, red=False, strike=False, bold=False):
    run = paragraph.add_run(text)
    run.font.strike = strike
    run.font.bold = bold
    if red:
        run.font.color.rgb = RED
    return run


def replace_with_runs(paragraph, chunks):
    paragraph.clear()
    for text, opts in chunks:
        add_run(paragraph, text, **opts)


def add_after(paragraph, text):
    new_p = paragraph.insert_paragraph_before("")
    # Move the newly inserted paragraph after the target paragraph.
    paragraph._p.addnext(new_p._p)
    add_run(new_p, text, red=True)
    return new_p


def revise():
    doc = Document(INPUT)

    # Title normalization.
    for p in doc.paragraphs:
        if p.text.strip() == "EduGenAI题目生成与质量评估平台设计文档":
            replace_with_runs(
                p,
                [
                    ("EduGenAI题目生成与质量评估平台设计文档", {"strike": True}),
                    ("EduGen AI：AI题目生成、质量评估与智能组卷平台设计文档", {"red": True, "bold": True}),
                ],
            )
            break

    # Remove the template reminder but keep it visible as a deletion.
    for p in doc.paragraphs:
        if "这一部分实际上容易和文档开始的表格" in p.text:
            replace_with_runs(p, [(p.text, {"strike": True})])
            break

    replacements = {
        "在此背景下，EduGenAI设计为一个面向教研场景的AI题目生成与质量评估平台，通过构建“AI生成+自动评分+人工审核+数据反馈+Prompt优化”的完整工作流，实现教育内容生产的系统化与可持续优化。": [
            ("在此背景下，EduGenAI设计为一个面向教研场景的AI题目生成与质量评估平台，通过构建“AI生成+自动评分+人工审核+数据反馈+Prompt优化”的完整工作流，实现教育内容生产的系统化与可持续优化。", {"strike": True}),
            ("在此背景下，EduGen AI设计为一个面向教研场景的AI题目生成、质量评估、人工审核与智能组卷平台，通过构建“真实AI生成+自动质量评分+人工审核+质量规则沉淀+智能组卷+数据反馈+Prompt优化”的完整工作流，实现教育内容生产的系统化、可追踪与可持续优化。", {"red": True}),
        ],
        "（7）扩展功能": [
            ("（7）扩展功能", {"strike": True}),
            ("（7）智能组卷与个性化测验", {"red": True, "bold": True}),
        ],
        "系统应预留题库管理与智能组卷等扩展能力。后续能够支持题目入库、知识点分类、智能抽题、试卷生成以及个性化练习推荐等功能，以满足更复杂的教育应用场景。": [
            ("系统应预留题库管理与智能组卷等扩展能力。后续能够支持题目入库、知识点分类、智能抽题、试卷生成以及个性化练习推荐等功能，以满足更复杂的教育应用场景。", {"strike": True}),
            ("系统已支持基于审核通过题目的组卷中心，教师可按学科、年级、知识点、题型、难度等条件筛选题目，也可通过题量、知识点、题型、难度和分层模板进行规则匹配式自动组卷。试卷可保存、继续编辑、查看，并通过浏览器打印/另存PDF；试卷题目以快照形式保存，支持试卷专属二次编辑，不影响原题库题目。", {"red": True}),
        ],
        "（4）题目管理": [
            ("（4）题目管理", {"strike": True}),
            ("（4）题目管理与二次修订", {"red": True, "bold": True}),
        ],
        "系统应支持对生成题目进行统一管理，包括题目的查看、编辑、删除、筛选与分类等功能。同时，系统应支持根据学科、知识点、难度等维度进行检索，便于后续题库建设与组卷使用。": [
            ("系统应支持对生成题目进行统一管理，包括题目的查看、编辑、删除、筛选与分类等功能。同时，系统应支持根据学科、知识点、难度等维度进行检索，便于后续题库建设与组卷使用。", {"strike": True}),
            ("系统支持对生成题目进行查看、编辑、删除、筛选与分类管理。教师可在审核页直接修订题干、选项、答案、解析、知识点和难度，并保留审核状态、问题标签与修改建议，使题目从AI初稿逐步转化为可用于组卷与教学评价的内容资产。", {"red": True}),
        ],
        "系统应支持人工审核流程，允许教师或教研人员对AI生成内容进行审核与修正。审核结果应包括“通过”“需修改”“不通过”等状态，并支持添加问题标签、审核备注与修改建议，以便后续问题归因与数据统计。": [
            ("系统应支持人工审核流程，允许教师或教研人员对AI生成内容进行审核与修正。审核结果应包括“通过”“需修改”“不通过”等状态，并支持添加问题标签、审核备注与修改建议，以便后续问题归因与数据统计。", {"strike": True}),
            ("系统支持人工审核流程，允许教师或教研人员对AI生成内容进行审核、修订与质量判断。审核结果包括“通过”“需修改”“不通过”“待审核”等状态，并支持添加问题标签、审核备注与修改建议。审核过程中形成的高频问题与判断标准可进一步沉淀为质量规则，进入质量规则库，用于后续Prompt优化和教研标准复用。", {"red": True}),
        ],
        "系统界面应简洁易用，操作流程应符合教师与教研人员的使用习惯，降低AI工具使用门槛。同时，系统需要支持 Mock 模式，在未配置真实 API 的情况下仍能够完成完整功能演示。": [
            ("系统界面应简洁易用，操作流程应符合教师与教研人员的使用习惯，降低AI工具使用门槛。同时，系统需要支持 Mock 模式，在未配置真实 API 的情况下仍能够完成完整功能演示。", {"strike": True}),
            ("系统界面应简洁易用，操作流程应符合教师与教研人员的使用习惯，降低AI工具使用门槛。当前版本采用真实OpenAI兼容接口作为题目生成与质量评估能力来源，用户需要通过环境变量配置OPENAI_API_KEY、OPENAI_BASE_URL和AI_MODEL后使用，避免演示数据与真实生成能力混淆。", {"red": True}),
        ],
        "系统应具备一定的异常处理能力。当AI接口调用失败或网络异常时，系统应自动进行错误提示或回退至Mock模式，避免功能完全不可用。同时，系统需要保证数据存储与读取的稳定性。": [
            ("系统应具备一定的异常处理能力。当AI接口调用失败或网络异常时，系统应自动进行错误提示或回退至Mock模式，避免功能完全不可用。同时，系统需要保证数据存储与读取的稳定性。", {"strike": True}),
            ("系统应具备清晰的异常处理能力。当AI接口调用失败、API Key缺失、模型名称错误或网络异常时，系统应直接提示错误原因，不再回退至Mock模式，从而保证生成结果均来自真实模型调用。同时，系统需要保证题目、审核记录、质量规则与试卷数据在本地存储中的稳定读写。", {"red": True}),
        ],
    }

    for p in doc.paragraphs:
        text = p.text.strip()
        if text in replacements:
            replace_with_runs(p, replacements[text])

    # Add quality rule library after artificial review requirement.
    for p in list(doc.paragraphs):
        if "系统支持人工审核流程，允许教师或教研人员对AI生成内容进行审核、修订与质量判断" in p.text:
            add_after(
                p,
                "（新增）质量规则库：系统应支持将人工审核过程中形成的问题标签、审核意见、修改建议和Prompt优化建议沉淀为结构化质量规则。规则字段包括规则名称、适用场景、关联知识点、问题标签、规则说明、改进建议和出现次数，帮助教师与AI产品经理持续积累可复用的教研经验。",
            )
            break

    # Add implementation-status note near design goals.
    for p in list(doc.paragraphs):
        if p.text.strip().startswith("为了弥补纯AI系统在教育场景中不可避免的误差问题"):
            add_after(
                p,
                "（新增）当前项目实现已从早期Mock演示版本升级为真实API配置版本，题目生成与质量评估均依赖真实OpenAI兼容接口；同时新增质量规则库、题目二次修订、智能组卷、试卷快照编辑与打印导出能力，使“AI生成—人工审核—规则沉淀—组卷应用”的闭环更加完整。",
            )
            break

    # Update market trend with implemented intelligent paper assembly.
    for p in doc.paragraphs:
        if p.text.strip().startswith("随着教育数字化进程加快"):
            replace_with_runs(
                p,
                [
                    (p.text, {"strike": True}),
                    ("随着教育数字化进程加快，题库资源、审核数据、质量规则与教学反馈等教育数据资产的重要性不断提升。未来AI教育产品之间的竞争，不仅取决于模型生成能力，还取决于数据积累、Prompt优化、质量治理与教研工作流设计能力。EduGen AI通过真实AI生成、人工审核、规则沉淀与智能组卷的结合，体现了AI教育产品从“内容生成工具”走向“教研生产系统”的发展方向。", {"red": True}),
                ],
            )
            break

    # Append a short change summary to the end.
    doc.add_paragraph("")
    summary = doc.add_paragraph()
    add_run(summary, "修订说明：", red=True, bold=True)
    summary_items = [
        "删除或修正了Mock模式、接口失败回退Mock等已不符合当前系统实现的描述。",
        "新增真实API必配、质量规则库、题目二次修订、智能组卷、个性化测验与试卷快照编辑等内容。",
        "保留原文结构，删除内容以删除线表示，新增内容以红色字体表示。",
    ]
    for item in summary_items:
        p = doc.add_paragraph(style=None)
        add_run(p, f"• {item}", red=True)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    revise()
